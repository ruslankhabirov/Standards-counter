from Architecture import MainAppenderClass
from Searcher import string_validator

import datetime
import os
from openpyxl import load_workbook
from openpyxl.styles import Alignment
from docx import Document
from Exceptions import IdIsEmpty


def connect_to_word_file():
    """Функция отвечает за создание объекта файла Word"""

    path = os.getcwd()
    files = os.listdir(path)

    for file in files:
        if file == "Идентификаторы.docx":
            document = Document("Идентификаторы.docx")
            return document

    document = Document()
    return document


def fill_word_file(document, name: str, serial_number: str, value: int, current_value: int, identificator):
    """Функция отвечает за корректное внесение идентификационных данных
    в Word-файл"""

    """Проверим, какой объект пришел нам в качестве переменной
    identificator - строка или словарь? Это необходимо для
    корректного ввода стандартов в таблицу в документе"""
    if isinstance(identificator, str):
        diversity = 0
    else:
        start = int(identificator[0][3:])
        end = int(identificator[1][3:])
        diversity = end - start
        identificator = identificator[0]

    table = document.add_table(rows=diversity+2, cols=5)
    row_of_first_column = table.cell(0, 0)
    row_of_second_column = table.cell(0, 1)
    row_of_third_column = table.cell(0, 2)
    row_of_fourth_column = table.cell(0, 3)
    row_of_fifth_column = table.cell(0, 4)

    row_of_first_column.text = "Имя стандарта"
    row_of_second_column.text = "Серийный номер"
    row_of_third_column.text = "Номинальный объём"
    row_of_fourth_column.text = "Действительный объём"
    row_of_fifth_column.text = "Идентификатор"

    if current_value == 0 or current_value == "0" or not current_value:
        current_value = ""

    for counter in range(0, diversity+1):
        first_column = table.cell(counter + 1, 0)
        second_column = table.cell(counter + 1, 1)
        third_column = table.cell(counter + 1, 2)
        fourth_column = table.cell(counter + 1, 3)
        fifth_column = table.cell(counter + 1, 4)

        first_column.text = name
        second_column.text = str(serial_number)
        third_column.text = str(value)
        fourth_column.text = current_value
        fifth_column.text = identificator[0:3] + str(int(identificator[3:]) + counter)


def find_all_projects():
    path = os.getcwd()
    files = os.listdir(path)

    projects_list = []

    for file_name in files:
        if "BCD" in file_name:
            projects_list.append(file_name[:-5])

    return sorted(projects_list, key=lambda x: x.split(" ")[1])


class AppenderClass(MainAppenderClass):
    """Конкретный класс, отвечающий за добавление новых стандартов
    в базу данных"""
    def __init__(self, project: str, catalogue_number: str, cas_number: str,
                 name_of_standard: str, manufacturer: str, serial_number: str,
                 nominal_value: int, number: int, minimum_balance: int, current_balance: int):
        self.project = project
        self.catalogue_number = catalogue_number
        self.cas_number = cas_number
        self.name_of_standard = name_of_standard
        self.manufacturer = manufacturer.upper()
        self.serial_number = serial_number
        self.nominal_value = int(nominal_value)
        self.number = int(number)
        self.minimum_balance = minimum_balance
        self.current_balance = current_balance

    def connection_to_db(self):
        """Функция находит и возвращает объект необходимой книги"""

        """Вначале получаем все имена файлов и папок, содержащихся в текущей
        директории"""
        path = os.getcwd()
        files = os.listdir(path)

        """Если с 1-й по 6-ой символы включительно совпадают с кодом выбранного проекта,
        то мы нашли нужный нам файл."""
        print(self.project[4:7])
        for file_name in files:
            if self.project[4:7] == file_name[4:7]:
                book = load_workbook(file_name)
                return book, file_name

    def get_id(self, book):
        """Функция, извлекающая ID из технической таблицы
        и положение последней строки из основной таблицы"""

        """Вначале получаем объект листа, внутрь которого будем вносить
        новые стандарты"""
        main_page_name = book.sheetnames[0]
        main_page = book[main_page_name]

        """Затем получаем объект листа, из которого будем извлекать текущий id"""
        id_page_name = book.sheetnames[1]
        id_page = book[id_page_name]

        """Проверим поле B1 на наличие значений. Если их нет, то
        возбуждаем ошибку"""
        if not id_page["B1"].value:
            raise IdIsEmpty("Значение поля идентификатора пусто!")

        """Если нет ошибок, то извлекаем необходимый ID
        и возвращаем объекты основного листа, технического листа
        и идентификатора"""
        identification = id_page["B1"].value

        """Также получаем объект каретки (последней заполненной
        строки)"""
        carriage = main_page.max_row

        return main_page, id_page, identification, carriage

    def get_previous_mass(self, page, carriage: int):
        """Функция осуществляет поиск по строкам таблицы, для нахождения последней
        общей суммы для вносимого стандарта"""

        """Проверим условие того, что клиент внёс CAS-номер"""
        if self.cas_number:
            for counter in range(carriage, 0, -1):
                current_cas_value = page["D" + str(counter)].value
                if current_cas_value == self.cas_number:
                    previous_mass = page["M" + str(counter)].value
                    if self.minimum_balance == 0 or self.minimum_balance == "0" or not self.minimum_balance:
                        self.minimum_balance = int(page["L" + str(counter)].value)
                    return previous_mass
            """Если CAS-номер есть, но ранее он не встречался,
            то тогда возвращаем нулевое значение"""
            return 0
        elif self.catalogue_number:
            """Если CAS-номера нет, осуществляем поиск по каталожнику"""
            for counter in range(carriage, 0, -1):
                current_catalogue_value = page["C" + str(counter)].value
                if current_catalogue_value == self.catalogue_number:
                    previous_mass = page["M" + str(counter)].value
                    if self.minimum_balance == 0 or self.minimum_balance == "0" or not self.minimum_balance:
                        self.minimum_balance = int(page["L" + str(counter)].value)
                    return previous_mass
            """Если каталожный номер есть, но ранее он не встречался,
            то тогда возвращаем нулевое значение"""
            return 0
        else:
            """В случае, если каталожный номер также отсутствует, приступаем
            к поиску по внесенному имени"""
            for counter in range(carriage, 0, -1):
                """Принято решение не сбрасывать в одну кучу одинаковые наименования,
                созданые разными производителями. Поэтому для стандартов без CAS-номера и
                каталожника существует дополнительное условие: совпадение не только по
                самому названию, но и по имени производителя"""
                current_name_value = page["E" + str(counter)].value
                current_manufacturer = page["F" + str(counter)].value
                if current_name_value == self.name_of_standard and current_manufacturer == self.manufacturer:
                    previous_mass = page["M" + str(counter)].value
                    if self.minimum_balance == 0 or self.minimum_balance == "0" or not self.minimum_balance:
                        self.minimum_balance = int(page["L" + str(counter)].value)
                    return previous_mass
            """Если стандарт совершенно уникален и никогда ранее не
            встречался, тогда возвращаем нулевое значение"""
            return 0

    def append_to_db(self, main_page, id_page, identification, carriage, previous_mass):
        """Функция, отвечающая за основной функционал добавления новых стандартов"""

        """Получаем новую строку"""
        new_row = carriage + 1

        """Получаем текущее время, которое будет внесено в таблицу как время
        занесения администратором стандарта"""
        now = datetime.datetime.now()

        """Итерируемся по количеству внесенных стандартов, заполняя поля ввода
        в новые строки"""
        end_of_id = ""

        """Получаем значение первого ID-номера, чтобы в дальнейшем добавить его
        в Word-отчёт"""
        first_id = identification[0:3] + str(int(identification[3:]) + 1)

        if self.minimum_balance == 0 or not self.minimum_balance:
            self.minimum_balance = 0

        for iteration in range(self.number):

            main_page["A" + str(new_row + iteration)].value = now.strftime("%d-%m-%Y %H:%M")
            main_page["B" + str(new_row + iteration)].value = "Администратор"
            main_page["C" + str(new_row + iteration)].value = self.catalogue_number
            main_page["D" + str(new_row + iteration)].value = self.cas_number
            main_page["E" + str(new_row + iteration)].value = self.name_of_standard
            main_page["F" + str(new_row + iteration)].value = self.manufacturer
            main_page["G" + str(new_row + iteration)].value = self.serial_number

            end_of_id = str(int(identification[3:]) + iteration + 1)
            main_page["H" + str(new_row + iteration)].value = identification[0:3] + end_of_id
            id_page["B1"].value = end_of_id

            string_validator(id_page["B1"].value)
            string_validator(main_page["H" + str(new_row + iteration)].value)

            main_page["I" + str(new_row + iteration)].value = self.nominal_value

            if self.current_balance == 0 or self.current_balance == "0" or not self.current_balance:
                main_page["K" + str(new_row + iteration)].value = self.nominal_value
                current_mass = int(int(previous_mass) + self.nominal_value * (iteration + 1))
            else:
                self.current_balance = int(self.current_balance)
                main_page["K" + str(new_row + iteration)].value = self.current_balance
                current_mass = int(int(previous_mass) + self.current_balance * (iteration + 1))

            main_page["M" + str(new_row + iteration)].value = current_mass

            main_page["L" + str(new_row + iteration)].value = self.minimum_balance

            main_page.row_dimensions[int(new_row + iteration)].height = 40
            for column in main_page[('A' + str(new_row + iteration)):('M' + str(new_row + iteration))][0]:
                column.alignment = Alignment(horizontal='center', vertical='center')

        id_page["B1"].value = identification[0:3] + end_of_id
        last_id = id_page["B1"].value

        if self.number == 1:
            return first_id
        else:
            return [first_id, last_id]
